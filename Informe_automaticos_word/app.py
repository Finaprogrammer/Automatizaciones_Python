from docx import Document

# Funciones para automatización de informes en Word


def crear_documento(nombre_archivo):
    """
    Crea un nuevo documento de Word y lo guarda con el nombre especificado.
    """
    doc = Document()
    doc.save(nombre_archivo)
    print(f"Documento '{nombre_archivo}' creado exitosamente.")

def agregar_titulo(nombre_archivo, titulo):
    """
    Agrega un título al documento de Word especificado.
    """
    doc = Document(nombre_archivo)
    doc.add_heading(titulo, level=1)
    doc.save(nombre_archivo)
    print(f"Título '{titulo}' agregado al documento '{nombre_archivo}'.")

def agregar_parrafo(nombre_archivo, texto):
    """
    Agrega un párrafo al documento de Word especificado.
    """
    doc = Document(nombre_archivo)
    doc.add_paragraph(texto)
    doc.save(nombre_archivo)
    print(f"Párrafo agregado al documento '{nombre_archivo}'.")

def agregar_tabla(nombre_archivo, datos):
    """
    Agrega una tabla al documento de Word especificado.
    """
    doc = Document(nombre_archivo)
    tabla = doc.add_table(rows=len(datos), cols=len(datos[0]))
    for i, fila in enumerate(datos):
        for j, celda in enumerate(fila):
            tabla.rows[i].cells[j].text = str(celda)
    doc.save(nombre_archivo)
    print(f"Tabla agregada al documento '{nombre_archivo}'.")

# Ejemplo de uso
if __name__ == "__main__":
    archivo = "informe.docx"
    crear_documento(archivo)
    agregar_titulo(archivo, "Informe Automático")
    agregar_parrafo(archivo, "Este es un informe generado automáticamente.")
    datos_tabla = [["Nombre", "Edad", "Ciudad"], ["Juan", 30, "Madrid"], ["Ana", 25, "Barcelona"]]
    agregar_tabla(archivo, datos_tabla)